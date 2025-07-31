from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import base64
import json
import os
from io import BytesIO
from models import Event
from docx.oxml.ns import qn

class ReportGenerator:
    def __init__(self):
        self.document = Document()
        
    def create_event_report(self, event):
        """Generate a Word document report for the given event"""
        
        # Set up document styles
        self._setup_styles()

        # Add event summary
        self._add_event_summary(event)
        
        # Add problems grouped by category
        self._add_problems_by_category(event)

        
        return self.document
    
    def _setup_styles(self):
        """Set up custom styles for the document"""
        styles = self.document.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        
        # Heading style
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        
        # Subheading style
        subheading_style = styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_style.font.name = 'Arial'
        subheading_style.font.size = Pt(14)
        subheading_style.font.bold = True

    def _add_event_summary(self, event):
        """Add event summary section"""

        # First three lines as one bold paragraph
        old_house_name = getattr(event, 'old_house_id', None)
        if old_house_name:
            applicant = f"申請人身份：{getattr(event, 'customer_name', '')}（{old_house_name}）"
        else:
            applicant = f"申請人身份：{getattr(event, 'customer_name', '')}"
        now = datetime.now()
        year = now.year
        month = now.month
        day = now.day
        check_date = f"檢查日期：{year} 年 {month} 月 {day} 日"
        title = f"{getattr(event.house, 'name', '')}{getattr(event, 'flat', '')}維修申請報告"
        # Combine all three lines
        first_three = f"{title}\n{applicant}\n{check_date}"
        first_paragraph = self.document.add_paragraph()
        run = first_paragraph.add_run(first_three)
        run.font.size = Pt(12)
        run.bold = True
        first_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set font for Chinese
        zh_font_name = 'Microsoft JhengHei'
        run.font.name = zh_font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font_name)

        # Fourth line: summary
        summary = (
            f"本人於上述日期對{getattr(event.house, 'name', '')}{getattr(event, 'flat', '')}進行全面檢查，"
            "發現單位存在多項建築結構、裝修質量及機電安裝方面的缺陷，部分情況屬嚴重，亟需安排維修處理。"
            "現提交詳細報告如下，懇請房委會安排相關部門盡快進行執漏及修復工作，以保障住戶安全及居住品質。"
        )
        summary_paragraph = self.document.add_paragraph(summary)
        summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in summary_paragraph.runs:
            run.font.name = zh_font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font_name)

    
    def _add_problems_by_category(self, event):
        """Rewrite: Group problems by important, then by category for non-important."""
        problems = event.problems if hasattr(event, 'problems') else []
        if not problems:
            self.document.add_paragraph('暫無問題記錄')
            return
        zh_font_name = 'Microsoft JhengHei'
        # Group problems
        important_problems = [p for p in problems if getattr(p, 'important', False)]
        other_problems = [p for p in problems if not getattr(p, 'important', False)]
        # 🛠️ 一、嚴重缺陷（需優先處理）
        if important_problems:
            imp_title = self.document.add_paragraph('🛠️ 一、嚴重缺陷（需優先處理）')
            imp_title.runs[0].font.bold = True
            imp_title.runs[0].font.name = zh_font_name
            imp_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font_name)
            for idx, problem in enumerate(important_problems, 1):
                self._add_problem_detail(problem, idx)
        # 🔧 二、其他需修復項目（按區域分類）
        if other_problems:
            oth_title = self.document.add_paragraph('🔧 二、其他需修復項目（按區域分類）')
            oth_title.runs[0].font.bold = True
            oth_title.runs[0].font.name = zh_font_name
            oth_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font_name)
            # Group by category
            categories = {}
            for p in other_problems:
                cat = getattr(p, 'category', '未分類')
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(p)
            for cat, cat_problems in categories.items():
                cat_title = self.document.add_paragraph(f' {cat} ')
                cat_title.runs[0].font.bold = True
                cat_title.runs[0].font.name = zh_font_name
                cat_title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font_name)
                for idx, problem in enumerate(cat_problems, 1):
                    self._add_problem_detail(problem, idx)

    def _add_problem_detail(self, problem, index):
        """Add detailed information for a single problem (one line description, then images)"""
        zh_font_name = 'Microsoft JhengHei'
        # Description as numbered item
        description = problem.description if hasattr(problem, 'description') else '無描述'
        desc_para = self.document.add_paragraph(f'{index}.\t{description}')
        for run in desc_para.runs:
            run.font.name = zh_font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), zh_font_name)
        # Images (if any)
        if hasattr(problem, 'image') and problem.image:
            images = problem.image
            if isinstance(images, str):
                try:
                    images = json.loads(images)
                except:
                    images = [images]
            elif not isinstance(images, list):
                images = [images]
            if images:
                for img_idx, img_data in enumerate(images):
                    try:
                        # Handle base64 image data
                        if isinstance(img_data, str) and ',' in img_data:
                            image_data = base64.b64decode(img_data.split(',')[1])
                        else:
                            image_data = base64.b64decode(img_data)
                        image_stream = BytesIO(image_data)
                        self.document.add_picture(image_stream, width=Inches(4))
                        # Add caption (optional, can be removed if not needed)
                        # caption = self.document.add_paragraph()
                        # caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # caption.add_run(f'圖片 {img_idx + 1}')
                    except Exception as e:
                        self.document.add_paragraph(f'圖片 {img_idx + 1} (無法載入)')
        self.document.add_paragraph()
    

    def save_document(self, filename):
        """Save the document to a file"""
        self.document.save(filename)
    
    def get_document_bytes(self):
        """Get the document as bytes for download"""
        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

def generate_event_report(event_url):
    """Generate a Word document report for the event with given URL"""
    event = Event.query.filter_by(url=event_url).first()
    if not event:
        return None
    
    generator = ReportGenerator()
    document = generator.create_event_report(event)
    
    return generator.get_document_bytes()

if __name__ == "__main__":
    # Test the report generator
    from app import app
    from models import db
    
    with app.app_context():
        event = Event.query.first()
        if event:
            generator = ReportGenerator()
            document = generator.create_event_report(event)
            generator.save_document('test_report.docx')
            print("Report generated successfully!")